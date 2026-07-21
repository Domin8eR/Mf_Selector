"""PDF and DOCX export for research reports.

PDF:   reportlab — title, header, bar chart (composite scores), top-funds table,
                   exec summary, fund narratives, key insights.
DOCX:  python-docx — same structure as text/tables (no chart image in DOCX).

Both files are written to the scratchpad directory and the path returned.
For local dev, the "path" is absolute; in production it would be an S3 key.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any


# ── Export directory (always local for MVP) ────────────────────────────────────
_EXPORT_DIR = Path("/private/tmp/altstreet_reports")
_EXPORT_DIR.mkdir(parents=True, exist_ok=True)


# ── PDF export ────────────────────────────────────────────────────────────────

def export_pdf(report: dict[str, Any]) -> str:
    """Generate a PDF for the report and return the file path."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.graphics.shapes import Drawing, Rect, String

    filename = _EXPORT_DIR / f"report_{report['id']}.pdf"
    doc = SimpleDocTemplate(
        str(filename),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    normal = styles["Normal"]
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=6,
    )
    h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=4,
    )
    small = ParagraphStyle("Small", parent=normal, fontSize=8, textColor=colors.grey)

    story: list[Any] = []

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph(report.get("title", "Research Report"), title_style))
    story.append(Paragraph(
        f"Category: {report.get('category', '—')} &nbsp;|&nbsp; "
        f"Snapshot: {report.get('snapshot_date', '—')} &nbsp;|&nbsp; "
        f"Rule: {report.get('rule_version_label', '—')}",
        small,
    ))
    story.append(Spacer(1, 6 * mm))

    # ── Bar chart — composite scores ──────────────────────────────────────────
    fund_narratives = report.get("fund_narratives", [])
    if fund_narratives:
        chart_w, chart_h = 460, 110
        drawing = Drawing(chart_w, chart_h)
        bar_w = min(60, (chart_w - 60) // len(fund_narratives))
        gap = 8
        max_score = max((f.get("composite_score", 0) or 0) for f in fund_narratives) or 100
        palette = [
            colors.HexColor("#2563EB"),
            colors.HexColor("#3B82F6"),
            colors.HexColor("#60A5FA"),
            colors.HexColor("#93C5FD"),
            colors.HexColor("#BFDBFE"),
        ]
        for i, fn in enumerate(fund_narratives[:5]):
            score = float(fn.get("composite_score", 0) or 0)
            bar_height = int((score / max_score) * 80)
            x = 40 + i * (bar_w + gap)
            rect = Rect(x, 15, bar_w, bar_height, fillColor=palette[i], strokeColor=None)
            drawing.add(rect)
            label = String(
                x + bar_w / 2, bar_height + 18,
                f"{score:.1f}",
                fontSize=7,
                textAnchor="middle",
            )
            drawing.add(label)
            rank_label = String(
                x + bar_w / 2, 4,
                f"#{fn.get('rank', i + 1)}",
                fontSize=7,
                textAnchor="middle",
                fillColor=colors.grey,
            )
            drawing.add(rank_label)

        story.append(Paragraph("Composite Scores — Top Ranked Funds", h2))
        story.append(drawing)
        story.append(Spacer(1, 4 * mm))

    # ── Top-funds table ───────────────────────────────────────────────────────
    story.append(Paragraph("Top-Ranked Research Candidates", h2))
    table_data = [["Rank", "Fund Name", "Score", "IR 3Y", "Sharpe", "6M Δ"]]
    for fn in fund_narratives:
        table_data.append([
            str(fn.get("rank", "")),
            _truncate(fn.get("fund_name", ""), 48),
            f"{float(fn.get('composite_score', 0) or 0):.2f}",
            f"{float(fn.get('information_ratio_3yr', 0) or 0):.4f}",
            f"{float(fn.get('sharpe_score', 0) or 0):.4f}",
            str(fn.get("rank_delta_6m", "—") or "—"),
        ])
    tbl = Table(table_data, colWidths=[15 * mm, 80 * mm, 20 * mm, 20 * mm, 20 * mm, 12 * mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("FONTSIZE", (0, 1), (-1, -1), 7),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E2E8F0")),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 4 * mm))

    # ── Executive summary ─────────────────────────────────────────────────────
    exec_summary = report.get("exec_summary", "")
    if exec_summary:
        story.append(Paragraph("Executive Summary", h2))
        for para in exec_summary.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), normal))
                story.append(Spacer(1, 2 * mm))

    # ── Fund narratives ───────────────────────────────────────────────────────
    story.append(Paragraph("Why Each Fund Stands Out", h2))
    for fn in fund_narratives:
        story.append(Paragraph(
            f"<b>#{fn.get('rank')} {fn.get('fund_name', '')}</b> — {fn.get('narrative', '')}",
            normal,
        ))
        story.append(Spacer(1, 2 * mm))

    # ── Key insights ──────────────────────────────────────────────────────────
    key_insights = report.get("key_insights", [])
    if key_insights:
        story.append(Paragraph("Key Insights", h2))
        for ki in key_insights:
            story.append(Paragraph(f"<b>{ki.get('heading', '')}</b>", normal))
            story.append(Paragraph(ki.get("body", ""), normal))
            story.append(Spacer(1, 2 * mm))

    # ── Footer disclaimer ─────────────────────────────────────────────────────
    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph(
        "This report is generated by the AltStreet research workspace. "
        "It presents metric-derived rankings and structural analysis only. "
        "It does not constitute investment advice, a buy/sell recommendation, "
        "or a suitability assessment for any individual or client.",
        small,
    ))

    doc.build(story)
    return str(filename)


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_docx(report: dict[str, Any]) -> str:
    """Generate a DOCX for the report and return the file path."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()

    # Title
    title_para = document.add_heading(report.get("title", "Research Report"), level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata line
    meta = document.add_paragraph(
        f"Category: {report.get('category', '—')}  |  "
        f"Snapshot: {report.get('snapshot_date', '—')}  |  "
        f"Rule: {report.get('rule_version_label', '—')}"
    )
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    document.add_paragraph()

    # Top-funds table
    document.add_heading("Top-Ranked Research Candidates", level=2)
    fund_narratives = report.get("fund_narratives", [])
    if fund_narratives:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(["Rank", "Fund Name", "Score", "IR 3Y", "Sharpe", "6M Δ"]):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True

        for fn in fund_narratives:
            row = table.add_row().cells
            row[0].text = str(fn.get("rank", ""))
            row[1].text = fn.get("fund_name", "")
            row[2].text = f"{float(fn.get('composite_score', 0) or 0):.2f}"
            row[3].text = f"{float(fn.get('information_ratio_3yr', 0) or 0):.4f}"
            row[4].text = f"{float(fn.get('sharpe_score', 0) or 0):.4f}"
            row[5].text = str(fn.get("rank_delta_6m", "—") or "—")

        document.add_paragraph()

    # Executive summary
    exec_summary = report.get("exec_summary", "")
    if exec_summary:
        document.add_heading("Executive Summary", level=2)
        for para in exec_summary.split("\n\n"):
            if para.strip():
                document.add_paragraph(para.strip())

    # Fund narratives
    document.add_heading("Why Each Fund Stands Out", level=2)
    for fn in fund_narratives:
        p = document.add_paragraph()
        run = p.add_run(f"#{fn.get('rank')} {fn.get('fund_name', '')} — ")
        run.bold = True
        p.add_run(fn.get("narrative", ""))

    # Key insights
    key_insights = report.get("key_insights", [])
    if key_insights:
        document.add_heading("Key Insights", level=2)
        for ki in key_insights:
            p = document.add_paragraph()
            run = p.add_run(ki.get("heading", "") + ": ")
            run.bold = True
            p.add_run(ki.get("body", ""))

    # Footer disclaimer
    document.add_paragraph()
    p = document.add_paragraph(
        "This report is generated by the AltStreet research workspace. "
        "It presents metric-derived rankings and structural analysis only. "
        "It does not constitute investment advice, a buy/sell recommendation, "
        "or a suitability assessment for any individual or client."
    )
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    filename = _EXPORT_DIR / f"report_{report['id']}.docx"
    document.save(str(filename))
    return str(filename)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _truncate(text: str, max_len: int) -> str:
    return text if len(text) <= max_len else text[: max_len - 1] + "…"
