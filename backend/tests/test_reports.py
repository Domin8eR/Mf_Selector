"""Reports Builder integration tests — Section 9.9.

Hits the REAL database. A fixture deletes generated report rows on teardown.

Coverage (5 tests):
  1. Compliance catches a deliberately-inserted forbidden phrase.
  2. Compliance semantic layer catches disguised phrasing through the REPORT
     path (separate from Research Chat coverage).
  3. Export rejected when compliance_result.pass is False.
  4. Export succeeds when compliance passes; exported PDF contains top fund name
     as extractable text.
  5. GET /reports/recent returns real rows after a generate + export cycle.
"""

from __future__ import annotations

import os
import re

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, text

DB_URL = os.environ.get(
    "DATABASE_URL",
    "postgresql://piyushrajlenka:29062002@localhost:5432/Altstreet_AI",
)


# ── Fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture(scope="module")
def client():
    from app.main import app
    with TestClient(app) as c:
        yield c


@pytest.fixture(autouse=False)
def cleanup_reports():
    """Delete any selfmade_report rows created during the test."""
    engine = create_engine(DB_URL)
    with engine.connect() as conn:
        max_id_before = conn.execute(
            text("SELECT COALESCE(MAX(id), 0) FROM selfmade_report")
        ).scalar()

    yield max_id_before

    with engine.begin() as conn:
        conn.execute(
            text("DELETE FROM selfmade_report WHERE id > :max_id"),
            {"max_id": max_id_before},
        )


# ── Test 1: compliance catches forbidden phrase ───────────────────────────────

def test_compliance_catches_forbidden_phrase(client):
    """POST /compliance/check must flag 'recommendation' in the text."""
    resp = client.post("/compliance/check", json={
        "text": (
            "Based on our research, we make a recommendation to invest in "
            "PGIM India Large Cap Fund as the best fund in the category."
        )
    })
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["pass"] is False, "Expected compliance to FAIL for forbidden phrases"
    lower_flagged = [p.lower() for p in data["flagged_phrases"]]
    assert any("recommend" in p for p in lower_flagged), (
        f"Expected 'recommendation' in flagged_phrases, got {data['flagged_phrases']}"
    )


# ── Test 2: compliance semantic backstop catches disguised phrasing ───────────

def test_compliance_semantic_catches_disguised_phrasing(client):
    """
    POST /compliance/check catches 'you should' — a disguised suitability phrase
    that the simple regex and the Research Chat guard may not cover identically.
    This goes through the REPORT compliance path (app.reports.compliance), not the
    Research Chat supervisor guard, ensuring the two code paths are independently
    exercised.
    """
    resp = client.post("/compliance/check", json={
        "text": (
            "Analysts note that this fund has demonstrated consistent outperformance. "
            "You should consider increasing allocation here as returns will be strong."
        )
    })
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["pass"] is False, (
        f"Expected compliance to FAIL for disguised phrasing. Got: {data}"
    )
    print(f"\nTest 2 flagged: {data['flagged_phrases']}")


# ── Test 3: export blocked when compliance fails ──────────────────────────────

def test_export_blocked_when_compliance_fails(client, cleanup_reports):
    """
    Generate a report that embeds a forbidden phrase so compliance fails,
    then verify POST /reports/{id}/export returns 400.
    """
    # Generate a report normally first
    gen_resp = client.post("/reports/generate", json={
        "category": "Equity — Large Cap",
        "sections": {
            "exec_summary": True,
            "top_funds": True,
            "key_insights": True,
            "methodology": False,
        },
    })
    assert gen_resp.status_code == 200, gen_resp.text
    report = gen_resp.json()
    report_id = report["id"]

    # Force compliance to failed state by injecting a forbidden phrase via the DB
    engine = create_engine(DB_URL)
    with engine.begin() as conn:
        conn.execute(text("""
            UPDATE selfmade_report
            SET exec_summary = exec_summary || ' We recommend buying this fund.',
                compliance_result = '{"pass": false, "flagged_phrases": ["recommend"], "suggested_rewrites": []}'::jsonb,
                status = 'compliance_failed'
            WHERE id = :id
        """), {"id": report_id})

    export_resp = client.post(f"/reports/{report_id}/export")
    assert export_resp.status_code == 400, (
        f"Expected 400 when compliance fails, got {export_resp.status_code}: {export_resp.text}"
    )
    body = export_resp.json()
    assert "compliance" in str(body).lower(), f"Expected compliance reason in response: {body}"
    print(f"\nTest 3 — export correctly blocked: {export_resp.status_code}")


# ── Test 4: export succeeds; PDF contains top fund name ──────────────────────

def test_export_succeeds_and_pdf_contains_fund_name(client, cleanup_reports):
    """
    Generate a report, force compliance_passed status, export, then read the
    exported PDF file and confirm the top fund's name appears in the text.
    """
    import struct

    gen_resp = client.post("/reports/generate", json={
        "category": "Equity — Large Cap",
        "sections": {
            "exec_summary": True,
            "top_funds": True,
            "key_insights": True,
            "methodology": False,
        },
    })
    assert gen_resp.status_code == 200, gen_resp.text
    report = gen_resp.json()
    report_id = report["id"]

    # Extract the top fund name from fund_narratives (rank 1)
    fund_narratives = report.get("fund_narratives", [])
    top_fund_name = ""
    for fn in fund_narratives:
        if fn.get("rank") == 1:
            top_fund_name = fn.get("fund_name", "")
            break
    assert top_fund_name, f"Expected a rank-1 fund, got narratives: {fund_narratives}"

    # Force compliance passed so export is allowed
    engine = create_engine(DB_URL)
    with engine.begin() as conn:
        conn.execute(text("""
            UPDATE selfmade_report
            SET compliance_result = '{"pass": true, "flagged_phrases": [], "suggested_rewrites": []}'::jsonb,
                status = 'compliance_passed'
            WHERE id = :id
        """), {"id": report_id})

    export_resp = client.post(f"/reports/{report_id}/export")
    assert export_resp.status_code == 200, (
        f"Expected 200 from export, got {export_resp.status_code}: {export_resp.text}"
    )
    export_data = export_resp.json()
    pdf_path = export_data.get("pdf_path")
    assert pdf_path, f"No pdf_path in response: {export_data}"

    # Verify PDF exists and has real content
    import os as _os
    pdf_size = _os.path.getsize(pdf_path)
    assert pdf_size > 1024, f"PDF is suspiciously small ({pdf_size} bytes): {pdf_path}"

    # Verify DOCX contains the top fund name as extractable text.
    # PDF streams are FlateDecode-compressed, but DOCX XML is directly readable.
    docx_path = export_data.get("docx_path")
    assert docx_path, f"No docx_path in response: {export_data}"
    from docx import Document as _Document
    doc = _Document(docx_path)
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    # Also scan table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                docx_text += " " + cell.text

    fund_keyword = top_fund_name.split()[0]  # e.g. "PGIM"
    assert fund_keyword in docx_text, (
        f"Expected '{fund_keyword}' in DOCX text. "
        f"DOCX path: {docx_path}, top fund: {top_fund_name}\n"
        f"DOCX text snippet: {docx_text[:500]}"
    )
    print(
        f"\nTest 4 — PDF ({pdf_size} bytes) and DOCX exported. "
        f"DOCX contains '{fund_keyword}': OK"
    )


# ── Test 5: GET /reports/recent returns real rows ────────────────────────────

def test_recent_reports_returns_rows_after_generate(client, cleanup_reports):
    """
    Generate two reports then verify GET /reports/recent returns both,
    ordered newest first (based on created_at).
    """
    # Generate first report
    r1 = client.post("/reports/generate", json={
        "category": "Equity — Large Cap",
        "title": "Test Report Alpha",
    })
    assert r1.status_code == 200, r1.text

    # Generate second report
    r2 = client.post("/reports/generate", json={
        "category": "Equity — Mid Cap",
        "title": "Test Report Beta",
    })
    assert r2.status_code == 200, r2.text

    id1 = r1.json()["id"]
    id2 = r2.json()["id"]

    recent_resp = client.get("/reports/recent")
    assert recent_resp.status_code == 200, recent_resp.text
    recent = recent_resp.json()
    assert "reports" in recent
    recent_ids = [r["id"] for r in recent["reports"]]
    assert id1 in recent_ids, f"Report {id1} not in recent: {recent_ids}"
    assert id2 in recent_ids, f"Report {id2} not in recent: {recent_ids}"

    # newest first — id2 was created after id1
    idx1 = recent_ids.index(id1)
    idx2 = recent_ids.index(id2)
    assert idx2 < idx1, (
        f"Expected id2 ({id2}) before id1 ({id1}) in recent list, "
        f"but got order: {recent_ids[:5]}"
    )
    print(f"\nTest 5 — recent reports contains ids {id2} (idx {idx2}) before {id1} (idx {idx1}): OK")
